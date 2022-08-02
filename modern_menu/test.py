from docx import Document
from docx.shared import Pt

document = Document('Entrees for Panda Garden.docx')

for table in document.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                print(paragraph.text)
                if 'ComboLoMein' in paragraph.text:
                    word = paragraph.text
                    p = document.add_paragraph("$13.25")

                    # Replace element with word
                    paragraph.text = paragraph.text.replace("ComboLoMein", p.text)

                    # Format/Style cell
                    run = cell.paragraphs[0].runs[0]
                    run.font.name = "Agency FB"
                    run.font.size = Pt(15)

document.save("Entrees for Panda Garden2.docx")